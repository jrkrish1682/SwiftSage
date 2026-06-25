"""
Assemble a Transformation Requirements Document (.docx) from
the mapping table and gap register produced by the Transformation Advisor.
"""
from __future__ import annotations

import io
from datetime import date
from typing import List

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from src.transformer.field_mapper import MappedField
from src.transformer.gap_analyzer import GapEntry

# ── Colour palette ─────────────────────────────────────────────────────────────
_NAVY   = RGBColor(0x1F, 0x4E, 0x79)
_BLUE   = RGBColor(0x2E, 0x75, 0xB6)
_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
_GRAY   = RGBColor(0x40, 0x40, 0x40)
_LIGHT  = RGBColor(0xD6, 0xE4, 0xF0)
_RED    = RGBColor(0xC0, 0x00, 0x00)
_AMBER  = RGBColor(0xBF, 0x89, 0x00)
_GREEN  = RGBColor(0x37, 0x86, 0x44)

_SEVERITY_COLOUR = {
    "CRITICAL": "C00000",
    "HIGH":     "BF8900",
    "MEDIUM":   "2E75B6",
    "LOW":      "378644",
}
_GAP_COLOUR = {
    "BLOCKING":    "C00000",
    "ENRICHMENT":  "BF8900",
    "CONDITIONAL": "2E75B6",
    "OUT_OF_SCOPE":"595959",
}
_MAP_COLOUR = {
    "DIRECT":   "378644",
    "DERIVED":  "BF8900",
    "SPLIT":    "2E75B6",
    "COMBINED": "7030A0",
    "UNMAPPED": "595959",
}


def _set_cell_bg(cell, hex_colour: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    tcPr.append(shd)


def _set_col_width(cell, inches: float) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _hcell(cell, text: str, width: float) -> None:
    _set_cell_bg(cell, "1F4E79")
    _set_col_width(cell, width)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(text)
    r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(10)
    r.font.color.rgb = _WHITE
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)


def _dcell(cell, text: str, width: float, shade: bool = False,
           colour: str | None = None, bold: bool = False) -> None:
    if shade:
        _set_cell_bg(cell, "D6E4F0")
    _set_col_width(cell, width)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.clear()
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(9); r.bold = bold
    if colour:
        r.font.color.rgb = RGBColor.from_string(colour)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)


def _heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.name  = "Calibri"
    run.font.color.rgb = _NAVY if level == 1 else _BLUE
    run.font.size  = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)


def _body(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11)
    r.italic = italic; r.font.color.rgb = _GRAY


def _bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(11)
        r.font.color.rgb = _NAVY
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.font.color.rgb = _GRAY


def generate_requirements_doc(
    mappings: List[MappedField],
    gaps: List[GapEntry],
    target_message_type: str,
    source_label: str = "Internal Bank Payment Message",
) -> bytes:
    """
    Generate a Transformation Requirements Document and return it as bytes
    suitable for st.download_button.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Footer ──────────────────────────────────────────────────────────────
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER; fp.clear()
    r = fp.add_run(
        f"SwiftSage  |  Transformation Requirements  |  {target_message_type}  |  Confidential"
    )
    r.font.name = "Calibri"; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = _GRAY

    # ── Cover ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(4)
    r = p.add_run("SwiftSage")
    r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(36)
    r.font.color.rgb = _NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Transformation Requirements Document")
    r.font.name = "Calibri"; r.font.size = Pt(18); r.font.color.rgb = _BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{source_label}  →  {target_message_type}")
    r.font.name = "Calibri"; r.font.size = Pt(13); r.italic = True
    r.font.color.rgb = _GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(20)
    r = p.add_run(f"Generated by SwiftSage  |  {date.today().isoformat()}")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = _GRAY

    doc.add_page_break()

    # ── 1. Executive Summary ─────────────────────────────────────────────────
    _heading(doc, "1.  Executive Summary", 1)

    direct   = [m for m in mappings if m.mapping_type == "DIRECT"]
    derived  = [m for m in mappings if m.mapping_type == "DERIVED"]
    split_   = [m for m in mappings if m.mapping_type == "SPLIT"]
    combined = [m for m in mappings if m.mapping_type == "COMBINED"]
    unmapped = [m for m in mappings if m.mapping_type == "UNMAPPED"]
    blocking = [g for g in gaps if g.gap_type == "BLOCKING" and not g.is_resolved]
    enrichment = [g for g in gaps if g.gap_type == "ENRICHMENT" and not g.is_resolved]
    conditional = [g for g in gaps if g.gap_type == "CONDITIONAL"]

    total_mapped = len(direct) + len(derived) + len(split_) + len(combined)
    complexity = (
        "HIGH"   if len(blocking) >= 3 or len(derived) >= 5 else
        "MEDIUM" if len(blocking) >= 1 or len(derived) >= 2 else
        "LOW"
    )

    _body(doc,
        f"This document describes the transformation requirements to map a {source_label} "
        f"to the ISO 20022 {target_message_type} (Customer Credit Transfer Initiation) standard. "
        f"It was generated by SwiftSage on {date.today().isoformat()}."
    )

    t = doc.add_table(rows=2, cols=4); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    ws = [1.5, 1.5, 1.5, 1.5]
    for i, (h, w) in enumerate(zip(
        ["Source Fields", "Fields Mapped", "Gaps Identified", "Migration Complexity"], ws
    )):
        _hcell(t.rows[0].cells[i], h, w)
    values = [str(len(mappings)), str(total_mapped), str(len(gaps)), complexity]
    colours = [None, None, None, {"HIGH": "C00000", "MEDIUM": "BF8900", "LOW": "378644"}[complexity]]
    for i, (v, c, w) in enumerate(zip(values, colours, ws)):
        _dcell(t.rows[1].cells[i], v, w, colour=c, bold=True)
    doc.add_paragraph()

    # ── 2. Field Mapping Table ────────────────────────────────────────────────
    _heading(doc, "2.  Field Mapping Table", 1)
    _body(doc,
        "Each internal field is mapped to its ISO 20022 equivalent. "
        "Mapping types: DIRECT (1:1), DERIVED (computed), SPLIT (1→many), "
        "COMBINED (many→1), UNMAPPED (no equivalent)."
    )

    col_w = [1.3, 1.4, 0.9, 1.8, 0.85, 2.55]
    headers = ["Source Field", "Sample Value", "Map Type", "ISO 20022 Target", "Confidence", "Business Rule / Notes"]
    t2 = doc.add_table(rows=len(mappings) + 1, cols=6)
    t2.style = "Table Grid"; t2.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (h, w) in enumerate(zip(headers, col_w)):
        _hcell(t2.rows[0].cells[i], h, w)

    for ri, m in enumerate(mappings, 1):
        shade = (ri % 2 == 0)
        map_colour = _MAP_COLOUR.get(m.mapping_type, "595959")
        _dcell(t2.rows[ri].cells[0], m.source_field, col_w[0], shade)
        _dcell(t2.rows[ri].cells[1], (m.source_value or "")[:40], col_w[1], shade)
        _dcell(t2.rows[ri].cells[2], m.mapping_type, col_w[2], shade, colour=map_colour, bold=True)
        target = m.iso20022_element or m.iso20022_xpath or "—"
        _dcell(t2.rows[ri].cells[3], target, col_w[3], shade)
        conf_col = {"HIGH": "378644", "MEDIUM": "BF8900", "LOW": "C00000"}.get(m.confidence, None)
        _dcell(t2.rows[ri].cells[4], m.confidence, col_w[4], shade, colour=conf_col)
        rule_text = m.business_rule
        if m.notes:
            rule_text += f" | {m.notes}"
        _dcell(t2.rows[ri].cells[5], rule_text[:200], col_w[5], shade)

    doc.add_paragraph()

    # ── 3. Gap Register ───────────────────────────────────────────────────────
    _heading(doc, "3.  Gap Register", 1)
    _body(doc,
        "Mandatory ISO 20022 fields that require resolution before the transformation "
        "can be considered complete. BLOCKING gaps will cause message rejection if not resolved."
    )

    open_gaps = [g for g in gaps if not g.is_resolved]
    if not open_gaps:
        _body(doc, "No open gaps identified. All mandatory fields have a source mapping.", italic=True)
    else:
        gw = [1.8, 2.0, 0.9, 0.85, 4.25]
        gheaders = ["ISO 20022 Field", "Business Label", "Gap Type", "Severity", "Recommended Resolution"]
        tg = doc.add_table(rows=len(open_gaps) + 1, cols=5)
        tg.style = "Table Grid"; tg.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (h, w) in enumerate(zip(gheaders, gw)):
            _hcell(tg.rows[0].cells[i], h, w)

        for ri, g in enumerate(open_gaps, 1):
            shade = (ri % 2 == 0)
            gap_colour = _GAP_COLOUR.get(g.gap_type, "595959")
            sev_colour = _SEVERITY_COLOUR.get(g.severity, None)
            _dcell(tg.rows[ri].cells[0], g.iso_xpath, gw[0], shade)
            _dcell(tg.rows[ri].cells[1], g.business_label, gw[1], shade)
            _dcell(tg.rows[ri].cells[2], g.gap_type, gw[2], shade, colour=gap_colour, bold=True)
            _dcell(tg.rows[ri].cells[3], g.severity, gw[3], shade, colour=sev_colour, bold=True)
            _dcell(tg.rows[ri].cells[4], g.recommendation, gw[4], shade)

    doc.add_paragraph()

    # ── 4. Unmapped Internal Fields ───────────────────────────────────────────
    if unmapped:
        _heading(doc, "4.  Unmapped Internal Fields", 1)
        _body(doc,
            "The following internal fields have no ISO 20022 equivalent. "
            "These are typically internal tracking, workflow, or audit fields "
            "that should be documented as out-of-scope for the transformation "
            "or handled via proprietary extensions."
        )
        for m in unmapped:
            _bullet(doc, f"{m.source_xpath} = '{m.source_value}' — {m.notes or 'No ISO 20022 equivalent.'}")
        doc.add_paragraph()

    # ── 5. Open Questions ──────────────────────────────────────────────────────
    _heading(doc, "5.  Open Questions for Business", 1)
    _body(doc, "The following decisions require Business Analyst or SME input before development can begin:")

    open_qs = [g for g in gaps if g.gap_type in ("BLOCKING", "CONDITIONAL") and not g.is_resolved]
    if open_qs:
        for i, g in enumerate(open_qs, 1):
            _bullet(doc,
                f"{g.recommendation}",
                bold_prefix=f"OQ-{i:02d} [{g.business_label}]: "
            )
    else:
        _body(doc, "No open questions — all ambiguous mappings have been resolved.", italic=True)

    doc.add_paragraph()

    # ── 6. Next Steps ─────────────────────────────────────────────────────────
    _heading(doc, "6.  Recommended Next Steps", 1)
    _bullet(doc, "Review all BLOCKING gaps with the ISO 20022 SME and resolve Open Questions above.",
            bold_prefix="Step 1 — Resolve Blocking Gaps: ")
    _bullet(doc, "Build or procure sort-code-to-BIC and sort-code-to-IBAN reference data services for DERIVED mappings.",
            bold_prefix="Step 2 — Reference Data: ")
    _bullet(doc, "Define business rules for conditional fields (ChrgBr, SvcLvl, LclInstrm) based on payment type.",
            bold_prefix="Step 3 — Business Rules: ")
    _bullet(doc, "Build the field-level transformation logic (XSLT, Java mapper, or middleware rule engine) based on the mapping table above.",
            bold_prefix="Step 4 — Implementation: ")
    _bullet(doc, "Test against the ISO 20022 XSD for pain.001.001.09 and validate with end-to-end test cases.",
            bold_prefix="Step 5 — Test & Validate: ")

    # ── Serialise to bytes ─────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
